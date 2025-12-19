import os
import io
import calendar
import json
from flask import Flask, render_template, redirect, url_for, request, flash, send_file, session
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from authlib.integrations.flask_client import OAuth
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Allow HTTP for local testing of Google Login
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-change-this'
app.config['DB_FILE'] = 'db.json'

# --- GOOGLE AUTH CONFIGURATION ---
# Replace with your actual credentials
app.config['GOOGLE_CLIENT_ID'] = 'YOUR_GOOGLE_CLIENT_ID_HERE'
app.config['GOOGLE_CLIENT_SECRET'] = 'YOUR_GOOGLE_CLIENT_SECRET_HERE'

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

oauth = OAuth(app)
google = oauth.register(
    name='google',
    client_id=app.config['GOOGLE_CLIENT_ID'],
    client_secret=app.config['GOOGLE_CLIENT_SECRET'],
    access_token_url='https://oauth2.googleapis.com/token',
    access_token_params=None,
    authorize_url='https://accounts.google.com/o/oauth2/auth',
    authorize_params=None,
    api_base_url='https://www.googleapis.com/oauth2/v1/',
    client_kwargs={'scope': 'openid email profile'},
)

# --- JSON DB HELPER FUNCTIONS ---
def init_db():
    if not os.path.exists(app.config['DB_FILE']):
        with open(app.config['DB_FILE'], 'w') as f:
            json.dump([], f)

def load_db():
    try:
        with open(app.config['DB_FILE'], 'r') as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return []

def save_db(users):
    with open(app.config['DB_FILE'], 'w') as f:
        json.dump(users, f, indent=4)

# Initialize DB on start
init_db()

# --- USER MODEL ---
class User(UserMixin):
    def __init__(self, id, username, password, is_superadmin=False, auth_type='local'):
        self.id = str(id)
        self.username = username
        self.password = password
        self.is_superadmin = is_superadmin
        self.auth_type = auth_type

    @staticmethod
    def get(user_id):
        users = load_db()
        for u in users:
            if str(u['id']) == str(user_id):
                return User(
                    id=u['id'], 
                    username=u['username'], 
                    password=u['password'],
                    is_superadmin=u.get('is_superadmin', False),
                    auth_type=u.get('auth_type', 'local')
                )
        return None

    @staticmethod
    def get_by_username(username):
        users = load_db()
        for u in users:
            if u['username'] == username:
                 return User(
                    id=u['id'], 
                    username=u['username'], 
                    password=u['password'],
                    is_superadmin=u.get('is_superadmin', False),
                    auth_type=u.get('auth_type', 'local')
                )
        return None

    @staticmethod
    def create(username, password, is_superadmin=False, auth_type='local'):
        users = load_db()
        # Auto-increment ID
        new_id = 1
        if users:
            new_id = max(u['id'] for u in users) + 1
            
        new_user_dict = {
            'id': new_id,
            'username': username,
            'password': password,
            'is_superadmin': is_superadmin,
            'auth_type': auth_type
        }
        users.append(new_user_dict)
        save_db(users)
        return User(new_id, username, password, is_superadmin, auth_type)

@login_manager.user_loader
def load_user(user_id):
    return User.get(user_id)

# --- DOCX GENERATION LOGIC ---
def generate_docx(target_year, uploaded_images):
    doc = Document()
    cal = calendar.Calendar(firstweekday=6)
    months = ["January", "February", "March", "April", "May", "June",
              "July", "August", "September", "October", "November", "December"]
    week_headers_short = ["S", "M", "T", "W", "T", "F", "S"]
    full_week_headers = ["Sunday", "Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday"]
    
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Mm(10)

    def remove_borders(table):
        tbl = table._tbl
        for cell in tbl.iter_tcs():
            tcPr = cell.tcPr
            tcBorders = OxmlElement('w:tcBorders')
            for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                edge = OxmlElement(f'w:{border}')
                edge.set(qn('w:val'), 'nil')
                tcBorders.append(edge)
            tcPr.append(tcBorders)

    # PAGE 1: YEAR SUMMARY
    title = doc.add_paragraph("World Wildlife Fund")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(22)
    title.runs[0].font.bold = True
    title.runs[0].font.name = "Arial"
    title.paragraph_format.space_after = Pt(0)

    subtitle = doc.add_paragraph(str(target_year))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.bold = True
    subtitle.runs[0].font.color.rgb = RGBColor(80, 80, 80)
    subtitle.paragraph_format.space_after = Pt(10)

    outer_table = doc.add_table(rows=4, cols=3)
    outer_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    outer_table.autofit = False 
    for col in outer_table.columns: col.width = Mm(63)
    remove_borders(outer_table)
    for row in outer_table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Mm(55)

    for i, month_name in enumerate(months):
        month_num = i + 1
        row_idx, col_idx = i // 3, i % 3
        cell = outer_table.cell(row_idx, col_idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        p = cell.add_paragraph(month_name.upper())
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.name = "Arial"
        
        month_weeks = cal.monthdayscalendar(target_year, month_num)
        inner_table = cell.add_table(rows=len(month_weeks) + 1, cols=7)
        inner_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        remove_borders(inner_table)
        
        hdr_row = inner_table.rows[0]
        hdr_row.height = Mm(3.5)
        for idx, day_char in enumerate(week_headers_short):
            c = hdr_row.cells[idx]
            c.text = day_char
            p_c = c.paragraphs[0]
            p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_c.runs[0]
            run.font.size = Pt(7)
            run.font.bold = True
            if idx == 0: run.font.color.rgb = RGBColor(255, 0, 0)
            else: run.font.color.rgb = RGBColor(100, 100, 100)

        for r, week in enumerate(month_weeks):
            row = inner_table.rows[r + 1]
            row.height = Mm(3.5)
            for c, day in enumerate(week):
                if day != 0:
                    row.cells[c].text = str(day)
                    p_day = row.cells[c].paragraphs[0]
                    p_day.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p_day.runs[0]
                    run.font.size = Pt(7)
                    if c == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 0, 0)

    doc.add_page_break()

    # PAGES 2-13
    for i, month_name in enumerate(months):
        month_num = i + 1
        
        img_key = f'image_{month_num}'
        if img_key in uploaded_images and uploaded_images[img_key]:
            try:
                img_stream = uploaded_images[img_key]
                img_stream.seek(0)
                doc.add_picture(img_stream, height=Mm(90))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.paragraphs[-1].paragraph_format.space_after = Pt(12)
            except:
                pass
        else:
            p = doc.add_paragraph(f"\n[No Image for {month_name}]\n")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        h = doc.add_heading(month_name.upper(), 1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.space_after = Pt(12)

        month_weeks = cal.monthdayscalendar(target_year, month_num)
        table = doc.add_table(rows=len(month_weeks) + 1, cols=7)
        table.style = "Table Grid"
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        hdr = table.rows[0]
        hdr.height = Mm(9)
        for idx, day in enumerate(full_week_headers):
            c = hdr.cells[idx]
            c.text = day
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.bold = True
            if idx == 0: p.runs[0].font.color.rgb = RGBColor(255, 0, 0)

        for r, week in enumerate(month_weeks):
            row = table.rows[r + 1]
            row.height = Mm(18)
            for c, day in enumerate(week):
                if day != 0:
                    cell = row.cells[c]
                    cell.text = str(day)
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.left_indent = Pt(4)
                    p.runs[0].font.size = Pt(14)
                    p.runs[0].font.bold = True
                    if c == 0: p.runs[0].font.color.rgb = RGBColor(255, 0, 0)

        if i < 11: doc.add_page_break()

    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f

# --- ROUTES ---

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
        
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        # Hardcoded Admin Check - Create in JSON if not exists
        if username == 'admin123' and password == 'admin123':
            user = User.get_by_username('admin123')
            if not user:
                hashed_pw = generate_password_hash('admin123', method='pbkdf2:sha256')
                user = User.create('admin123', hashed_pw, is_superadmin=True)
            login_user(user)
            return redirect(url_for('superadmin'))

        # Standard User Login
        user = User.get_by_username(username)
        if user and user.password and check_password_hash(user.password, password):
            login_user(user)
            return redirect(url_for('dashboard'))
        else:
            flash('Login failed. Check details.')
            
    return render_template('login.html')

@app.route('/google_login')
def google_login():
    return google.authorize_redirect(url_for('google_callback', _external=True))

@app.route('/google_callback')
def google_callback():
    resp = google.authorize_access_token()
    user_info = google.get('userinfo').json()
    email = user_info['email']
    
    user = User.get_by_username(email)
    if not user:
        # Create Google User (No password)
        user = User.create(username=email, password=None, is_superadmin=False, auth_type='google')
    
    login_user(user)
    return redirect(url_for('dashboard'))

@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))

    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        if User.get_by_username(username):
            flash('Username already exists.')
            return redirect(url_for('register'))
            
        hashed_pw = generate_password_hash(password, method='pbkdf2:sha256')
        new_user = User.create(username, hashed_pw, is_superadmin=False)
        
        login_user(new_user)
        return redirect(url_for('dashboard'))
        
    return render_template('register.html')

@app.route('/dashboard', methods=['GET', 'POST'])
@login_required
def dashboard():
    if current_user.is_superadmin:
        return redirect(url_for('superadmin'))
    return render_template('dashboard.html')

@app.route('/download', methods=['POST'])
@login_required
def download_calendar():
    year = int(request.form.get('year'))
    uploaded_images = {}
    for i in range(1, 13):
        file = request.files.get(f'image_{i}')
        if file and file.filename != '':
            in_memory_file = io.BytesIO()
            file.save(in_memory_file)
            uploaded_images[f'image_{i}'] = in_memory_file
    
    file_stream = generate_docx(year, uploaded_images)
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f'WWF_Calendar_{year}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/superadmin')
@login_required
def superadmin():
    if not current_user.is_superadmin:
        return redirect(url_for('dashboard'))
    # Load all users directly from JSON to pass to template
    users_data = load_db()
    return render_template('superadmin.html', users=users_data)

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True)